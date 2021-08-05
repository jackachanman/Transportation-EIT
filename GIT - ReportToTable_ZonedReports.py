# -*- coding: utf-8 -*-
"""
Created on Thu Jan  2 12:42:29 2020

@author: 20JCC
"""
import re, os, glob, csv
path = input('Enter directory where zoned reports are located: ')
import pandas as pd
from pandas import DataFrame
import decimal
from decimal import Decimal
result_type = input("record 'all' or 'critical' results? (pressing 'enter' defaults to 'all'):  ")
if result_type == "critical":
    vc_crit = float(input("Please input minimum threshold for 'critical' v/c ratio: "))
    los_crit = input("Please input threshold for 'critical' level of service ('A', 'B', 'C', 'D', 'E', or 'F'): ")
else:
    pass
#from collections import namedtuple
# %% FUNCTIONS
def is_number(s):
    try:
        float(s)
        return True
    except ValueError:
        return False
def Sub_DataFrame(df):
    subdf = df.iloc[5:, 2:]                                                     # extracting sub dataframe, extracting row 5 and on because 'Sign Control' is in different rows for TWSC and AWSC
    subdf_ind = []
    for string in list(df.iloc[5:, 0]):
        subdf_ind.append(string.strip())                                        # stripping leading and trailing spaces
    subdf.columns = df.iloc[4, 2:]
    subdf.index = subdf_ind
    return subdf
def is_TWSC(df):
    df = Sub_DataFrame(df)
    if df.loc['Sign Control'].str.count('Stop').sum() <= 2:                     # location of 'Sign control'
#        print (name + ' is a Two-way stopped controlled (TWSC)' )
        return True
    elif df.loc['Sign Control'].str.count('Stop').sum() == 3 or 4:
#        print (name + ' is All-way stopped controlled (AWSC)')
        return False
def Stopped_Approach(df):
    data_df = Sub_DataFrame(df)
    # from https://stackoverflow.com/questions/41289074/find-column-names-when-row-element-meets-a-criteria-pandas
    stopped_approach = ((data_df.loc['Sign Control'] == 'Stop') | (data_df.loc['Sign Control'] == 'Yield'))[lambda x: x].index.tolist()
    for g in range(len(stopped_approach)):
        if stopped_approach[g][0] == 'W':
            stopped_approach[g] = 'W'
        elif stopped_approach[g][0] == 'E':
            stopped_approach[g] = 'E'
        elif stopped_approach[g][0] == 'N':
            stopped_approach[g] = 'N'
        elif stopped_approach[g][0] == 'S':
            stopped_approach[g] = 'S'
    return stopped_approach
def Stopped_Movements(df, stopped_approach):                                    # finds stop movements. I believe this covers all scenarios
    pseudo_stpd_mvmt = []
    stopped_movement = []
    data_df = Sub_DataFrame(df)
    lane_grp = ((data_df.loc['Lanes']!= "") & (data_df.loc['Lanes']!= "0"))[lambda x: x].index.tolist()        # finds lane groups
    if is_TWSC(df) == True:
        r_flare = (data_df.loc['Right turn flare (veh)'] != '')[lambda x: x].index.tolist()                    # find movements that are right turn flares.
        if 'r_flare' in locals():
            for mvmt in r_flare:
                print (mvmt + ' is a right turn flare')
            for mvmt in lane_grp:
                if mvmt[0] in stopped_approach:                                 # find stopped movements
        #            print( mvmt + ' is stopped movement')
                    pseudo_stpd_mvmt.append(mvmt)
            for mvmt in lane_grp:                                               # stopped movments that are right turn flares are not recorded SEPARATELY in 'Direction, Lane #'
                if (mvmt in pseudo_stpd_mvmt) and (mvmt not in r_flare):
                    print (mvmt + ' is ....')
                    stopped_movement.append(mvmt)
        else:
            for mvmt in lane_grp:
                if mvmt[0] in stopped_approach:
                    stopped_movement.append(mvmt)
    elif is_TWSC(df) == False:                                                  # all way stopped approach does not have right turn flare rows? - confirmed with Uni-Fab AWSC - fake. This is a weird characteristic of HCM 2000 Unsig reporting
        for mvmt in lane_grp:
            if mvmt[0] in stopped_approach:
                stopped_movement.append(mvmt)
    return stopped_movement
#    return stopped_movement, r_flare, lane_grp
    # return stopped_movement, lane_grp
def HCM_unsig_result(df):
    df = Sub_DataFrame(df)                                                      # calls previously defined function, since the Sub_DataFrame cleans the index trailing/leading spaces
    results = df.loc['Direction, Lane #':'Approach LOS', :]                     # extract initial sub-dataframe
    col = []
    for item in df.loc['Direction, Lane #']:                                    # get names of Direction that is not ' ', will be used as new column headers
        if item != '':
            col.append(item)
    results = results.iloc[results.index.get_loc('Direction, Lane #')+1:\
                           results.index.get_loc('Approach LOS')+1, :len(col)]  # extract sub-dataframe based on integer index
    results.columns = col                                                       # declare new column header
    lane_fig = []                                                               # lane_fig is in format: <1, 2>, <1>, etc.
    for item in df.loc['Lanes',:].tolist():
        if (item!="") and (item!="0"):
            lane_fig.append(item)
    return results, lane_fig                                                    # lane_fig is used as input for Unsig_Report_Table function.
def Intersection_Summ(df):
    sub_df = Sub_DataFrame(df)
    intersection_summ = sub_df.loc['Intersection Summary':, ]
    if df.iloc[0, 0] == 'HCM Unsignalized Intersection Capacity Analysis':
        for item in intersection_summ.loc['Delay']:
            if item != '':
                intersection_delay = item
        for item in intersection_summ.loc['Level of Service']:
            if item != '':
                intersection_LOS = item
    elif df.iloc[0, 0] == 'Lanes, Volumes, Timings':
        indices_delay = [c for c, d in enumerate(intersection_summ.index.tolist()) if 'Intersection Signal Delay' in d] # returns indices of Intersection delay
        intersection_delay = intersection_summ.index.tolist()[indices_delay[0]].split(': ')[1]
        indices_LOS = [b for b, a in enumerate(intersection_summ.iloc[indices_delay[0]]) if 'Intersection LOS' in a]    # intersection LOS is on the same row as intersection delay
        intersection_LOS = intersection_summ.iloc[indices_delay[0], indices_LOS[0]].split(': ')[1]
    return intersection_delay, intersection_LOS
# unsignalized - LOS Defintions
def Unsignalized_LOS(delay):
    if delay <= 10:
        los = 'A'
    elif 10 < delay <= 15:
        los = 'B'
    elif 15 < delay <= 25:
        los = 'C'
    elif 25 < delay <= 35:
        los = 'D'
    elif 35 < delay <= 50:
        los = 'E'
    elif 50 < delay:
        los = 'F'
    return los
# signalized - LOS Defintions
def Signalized_LOS(delay):
    if delay <= 10:
        los = 'A'
    elif 10 < delay <= 20:
        los = 'B'
    elif 20 < delay <= 35:
        los = 'C'
    elif 35 < delay <= 55:
        los = 'D'
    elif 55 < delay <= 80:
        los = 'E'
    elif 80 < delay:
        los = 'F'
    return los
def Unsig_Report_Table(result_df, stopped_result, column_name, stopped_movement, lane_fig):
    table_mvmt = []
    for mvmt in stopped_movement:
        if mvmt[2] == 'R':
            table_mvmt.append(mvmt[0:2] + ' right')
        elif mvmt[2] == 'T':
            table_mvmt.append(mvmt[0:2] + ' through')
        elif mvmt[2] == 'L':
            table_mvmt.append(mvmt[0:2] + ' left')
        lane_arr = []                                                           # lane_arr is lane configuration in format: EBT, EBTR, EBTL, EBTRL. this works because '<' and '>' applies to through movement only.... i believe.
        for i , k in zip(stopped_movement, lane_fig):
            if '<' in k and '>' not in k:
                i += 'L'
            elif '>' in k and '<' not in k:
                i += 'R'
            elif '<' in k and '>' in k:
                i += 'LR'
            else:
                next
            lane_arr.append(i)
    table = result_df.T                                                         # transpose
    table = table.loc[table.index.isin(stopped_result)]                         # remove indices that are not stopped approaches
    table = table[table.columns.intersection(column_name)]                      # removing caolumns that are not relevant results
    table = table.reindex(columns = column_name)                                # reorders columns
    # table.index = table_mvmt                                                    # replace indices
    table.index = lane_arr
    return table
def Sig_Report_Table(df, column_name):
    data_df = Sub_DataFrame(df)
    lane_grp = ((data_df.loc['Lane Configurations'] != "") & (data_df.loc['Lane Configurations'] != "0"))[lambda x: x].index.tolist()        # finds lane groups
    table = data_df.T                                                           # transpose
    table = table.loc[table.index.isin(lane_grp)]                               # remove indices that are not stopped approaches
    table = table[table.columns.intersection(column_name)]                      # removing caolumns that are not relevant results
    table = table.reindex(columns = column_name)                                # reorders columns
    table_mvmt = []
    for mvmt in lane_grp:
        if mvmt[2] == 'R':
            table_mvmt.append(mvmt[0:2] + ' right')
        elif mvmt[2] == 'T':
            table_mvmt.append(mvmt[0:2] + ' through')
        elif mvmt[2] == 'L':
            table_mvmt.append(mvmt[0:2] + ' left')      
    lane_fig = []                                                               # lane_fig is in format: <1, 2>, <1>, etc.
    lane_arr = []                                                               # lane_arr is lane configuration in format: EBT, EBTR, EBTL, EBTRL. this works because '<' and '>' applies to through movement only.... i believe.
    for item in data_df.loc['Lane Configurations',:].tolist():
        if (item!="") and (item!="0"):
            lane_fig.append(item)
    for i , k in zip(lane_grp, lane_fig):
        if '<' in k and '>' not in k:
            i += 'L'
        elif '>' in k and '<' not in k:
            i += 'R'
        elif '<' in k and '>' in k:
            i += 'LR'
        else:
            next
        lane_arr.append(i)
    # table.index = table_mvmt                                                    # replace indices
    table.index = lane_arr
    return table
def Name_Shorten(string, acrynm_dict):                                          # case sensitive
    for key in acrynm_dict.keys():
        string = string.replace(key,acrynm_dict[key])
        if len(string) > 31:
            string = string.replace(' ','')
            if len(string) > 31:
                string = string[0:31]
            else:
                pass
        else:
            pass
    return string
def Table_Filter(df,vc,los, colname):
    loscol = colname[1]
    vccol = colname[0]
    # stripping defacto left turn lane 'dl' and right turn lane 'dr' from v/c result. See page 10-15 of Synchro 10 guide
    #https://stackoverflow.com/questions/13682044/remove-unwanted-parts-from-strings-in-a-column
    df[vccol] = df[vccol].map(lambda x: x.rstrip('dl'))         
    df[vccol] = df[vccol].map(lambda x: x.rstrip('dr'))
    if los == 'A':
        df = df.loc[(df[loscol] == 'A' )\
                                |(df[loscol] == 'B')\
                                |(df[loscol] == 'C')\
                                |(df[loscol] == 'D')\
                                |(df[loscol] == 'E')\
                                |(df[loscol] == 'F')\
                                |(pd.to_numeric(df[vccol], downcast='float')>=vc)]
    elif los == 'B':
        df = df.loc[(df[loscol] == 'B')\
                                |(df[loscol] == 'C')\
                                |(df[loscol] == 'D')\
                                |(df[loscol] == 'E')\
                                |(df[loscol] == 'F')\
                                |(pd.to_numeric(df[vccol], downcast='float')>=vc)]
    elif los == 'C':
        df = df.loc[(df[loscol] == 'C')\
                                |(df[loscol] == 'D')\
                                |(df[loscol] == 'E')\
                                |(df[loscol] == 'F')\
                                |(pd.to_numeric(df[vccol], downcast='float')>=vc)]
    elif los == 'D':
        df = df.loc[(df[loscol] == 'D')\
                                |(df[loscol] == 'E')\
                                |(df[loscol] == 'F')\
                                |(pd.to_numeric(df[vccol], downcast='float')>=vc)]
    elif los == 'E':
        df = df.loc[(df[loscol] == 'E')\
                                |(df[loscol] == 'F')\
                                |(pd.to_numeric(df[vccol], downcast='float')>=vc)]
    elif los == 'F':
        df = df.loc[(df[loscol] == 'F')
                                |(pd.to_numeric(df[vccol], downcast='float')>=vc)]
    return df

# %% Read Data
#create named tuple
#Intersection = namedtuple('Intersection',['name', 'peak', 'scenario', 'report'])
os.chdir(path)
dict_csv = {}
dict_table = {}
dict_table_filtered = {}
int_name = []
int_time = []
int_scenario = []
int_key = []
int_node = []
int_nodekey = []
for txtfile in sorted(glob.glob("*.txt")):
    f = []                                                                      # reset list after each text file is converted to dataframe
    with open (txtfile) as csvfile:
        csv_reader = csv.reader(csvfile, delimiter = '\t')
        for row in csv_reader:
            f.append(row)
    data_raw = DataFrame(f).fillna('')
    sep = []
    for index, row in data_raw.iterrows():
        if (row[0] == 'HCM Unsignalized Intersection Capacity Analysis') or (row[0] == 'Lanes, Volumes, Timings'):
            sep.append(int(index))
    sep.append(len(data_raw))
    for i in range(len(sep)-1):
        data = data_raw.iloc[sep[i]:sep[i+1], :]
        #intersection name
        id_text = data.iloc[1, 0]
        id_num, name = id_text.split(': ')
        #intersection peak hour
        peak = data.iloc[0, 1]
        #intersection scenario
        scen = data.iloc[1, 1]
        #intersection type
        report = data.iloc[0, 0]
        key = (name, peak, scen, report, id_num)
        # dictioanry of dataframes with tuples as keys
        dict_csv[key] = data.copy()
        int_name.append(name)
        int_time.append(peak)
        int_scenario.append(scen)
        int_key.append(key)
        int_node.append(id_num)
        int_nodekey.append((id_num,name))
# =============================================================================
# # Extracting relevant results from results_df and puttin into table"""
# =============================================================================      
        if report == 'HCM Unsignalized Intersection Capacity Analysis':
            stpd_aprch = Stopped_Approach(data)
            stpd_mvmt = Stopped_Movements(data, stpd_aprch)
            """ Translates to EB 1 EB 2 WB 1 WB 2 NB 1 NB 2 etc. """
            stpd_rslt = []                                                      # this works cuz HCm 2000 Unsig limits to 2 lanes per approach
            for i in stpd_aprch:
                count = 0
                for j in stpd_mvmt:
                    if i == j[0]:
                        count += 1
                        stpd_rslt.append(j[0:2] + ' ' + str(count))
            #check that stpd_rslt and stpd_mvmt are the same length
            if len(stpd_rslt) == len(stpd_mvmt):
                print ('stpd_rslt and stpd_mvmt are the same length << AWESOME')
            else:
                print ('stpd_rslt and stpd_mvmt are NOT the same length << BAD')
            if is_TWSC(data) == True:
                table_columns = ['Volume to Capacity','Lane LOS', 'Control Delay (s)','Queue Length 95th (m)'] # for TWSC ... Control delay or approach delay
                table_df = Unsig_Report_Table(HCM_unsig_result(data)[0], stpd_rslt, table_columns, stpd_mvmt, HCM_unsig_result(data)[1])
                # finds missing LOS and fills in
                for index, _ in table_df.iterrows():
                    if (table_df.loc[index, 'Lane LOS'] == '') and (table_df.loc[index, 'Control Delay (s)'] != ''):
                        table_df.loc[index,'Lane LOS'] = Unsignalized_LOS(float(table_df.loc[index, 'Control Delay (s)']))
                    else:
                        pass
                    # modify results if 'Err' is given
                    table_df.loc[table_df['Control Delay (s)'] == 'Err', 'Control Delay (s)'] = '>200' # added for 20-3410 Rymal Road @ Arcadia and Arrowhead
                    # rounds to integer if delay is over 100
                    try:
                        if float(table_df.loc[index,'Control Delay (s)']) >= 100:
                            table_df.loc[index, 'Control Delay (s)'] = str(Decimal(float(table_df.loc[index, 'Control Delay (s)'])).quantize(Decimal('1'),rounding = decimal.ROUND_HALF_UP))
                        else:
                            pass
                    except:
                        pass
                    # extracts just the number from queue and rounds to zero decimal place
                    temp = re.findall("\d+\.\d+",table_df.loc[index, 'Queue Length 95th (m)'])  # extracts just the number from queue. creates a .  From https://stackoverflow.com/questions/4703390/how-to-extract-a-floating-number-from-a-string
                    if temp:                                                                    # if temp is not empty
                        table_df.loc[index, 'Queue Length 95th (m)'] = Decimal(float(temp[0])).quantize(Decimal('1'),rounding = decimal.ROUND_HALF_UP)
                    elif not temp:                                                              # if temp is empty
                        pass
                    # filters for critical movements
                    if result_type == 'critical':
                        table_df_filtered = Table_Filter(table_df, vc_crit, los_crit, table_columns)
                    else:
                        pass
            elif is_TWSC(data) == False:
                table_columns = ['Degree Utilization, x', 'Approach LOS','Control Delay (s)']
                table_df = Unsig_Report_Table(HCM_unsig_result(data)[0], stpd_rslt, table_columns, stpd_mvmt, HCM_unsig_result(data)[1])
                if result_type == 'critical':
                    table_df_filtered = Table_Filter(table_df.copy(), vc_crit, los_crit, table_columns)
                else:
                    pass          
                # finds missing LOS and fills in
                for index, _ in table_df.iterrows():
                    if (table_df.loc[index, 'Approach LOS'] == '') and (table_df.loc[index, 'Control Delay (s)'] != ''):
                        table_df.loc[index, 'Approach LOS'] = Unsignalized_LOS(float(table_df.loc[index, 'Control Delay (s)']))
                    else:
                        pass
                    # modifies results if 'Err' is given
                    table_df.loc[table_df['Control Delay (s)'] == 'Err', 'Control Delay (s)'] = '>200' # added for 20-3410 Rymal Road @ Arcadia and Arrowhead
                    # rounds to integer if delay is over 100
                    try:
                        if float(table_df.loc[index,'Control Delay (s)']) >= 100:
                            table_df.loc[index, 'Control Delay (s)'] = str(Decimal(float(table_df.loc[index, 'Control Delay (s)'])).quantize(Decimal('1'),rounding = decimal.ROUND_HALF_UP))
                        else:
                            pass
                    except:
                        pass
                # AWSC does not report queue
                # intersection summary
                int_delay, int_LOS = Intersection_Summ(data)
                
                # filters for critical movments
                if result_type == 'critical':
                    if table_df_filtered.empty:
                        pass
                    else:
                        table_df_filtered.loc['Overall', 'Control Delay (s)'] = int_delay
                        table_df_filtered.loc['Overall', 'Approach LOS'] = int_LOS
                        table_df_filtered = table_df_filtered.fillna('-')
                else:
                    pass
                table_df.loc['Overall', 'Control Delay (s)'] = int_delay
                table_df.loc['Overall', 'Approach LOS'] = int_LOS
                table_df = table_df.fillna('-')
        elif report == 'Lanes, Volumes, Timings': 
            table_columns = ['v/c Ratio', 'LOS', 'Total Delay', 'Queue Length 95th (m)']    # rows from textfile that we want
            table_df = Sig_Report_Table(data, table_columns)
            if result_type == 'critical':
                table_df_filtered = Table_Filter(table_df.copy(), vc_crit, los_crit, table_columns)
            else:
                pass   
            # finds missing LOS and fills in
            for index, _ in table_df.iterrows():
                if (table_df.loc[index, 'LOS'] == '') and (table_df.loc[index, 'Total Delay'] != ''):
                    table_df.loc[index, 'LOS'] = Signalized_LOS(float(table_df.loc[index, 'Total Delay']))
                else:
                    pass
                # modfies results if 'Err' is given
                table_df.loc[table_df['Total Delay'] == 'Err', 'Total Delay'] = '>1000'
                try:
                    # rounds to integer if delay is over 100
                    if (table_df.loc[index,'Total Delay'] != '') and (float(table_df.loc[index,'Total Delay']) >= 100):
                        table_df.loc[index, 'Total Delay'] = str(Decimal(float(table_df.loc[index, 'Total Delay'])).quantize(Decimal('1'),rounding = decimal.ROUND_HALF_UP))
                    else:
                        pass
                except:
                    pass
                # extracts just the number from queue and rounds to zero decimal place
                temp = re.findall("\d+\.\d+", table_df.loc[index, 'Queue Length 95th (m)']) # extracts just the number from queue. creates a .  From https://stackoverflow.com/questions/4703390/how-to-extract-a-floating-number-from-a-string
                if temp:                                                                    # if temp is not empty
                    table_df.loc[index, 'Queue Length 95th (m)'] = Decimal(float(temp[0])).quantize(Decimal('1'),rounding = decimal.ROUND_HALF_UP)
                elif not temp:                                                              # if temp is empty
                    pass
                int_delay, int_LOS = Intersection_Summ(data)
                #Filters critical movements
                if result_type == 'critical':
                    if table_df_filtered.empty:                                 # don't bother adding overall intersection results if there are no critical movements
                        pass
                    else:
                        table_df_filtered.loc['Overall', 'Total Delay'] = int_delay
                        table_df_filtered.loc['Overall', 'LOS'] = int_LOS
                        table_df_filtered = table_df_filtered.fillna('-')
                else:
                    pass
                table_df.loc['Overall', 'Total Delay'] = int_delay
                table_df.loc['Overall', 'LOS'] = int_LOS
                table_df = table_df.fillna('-')
                
        dict_table[key] = table_df.copy()
        if result_type == 'critical':
            dict_table_filtered[key] = table_df_filtered.copy()
        else:
            pass

"""
^^^^^^^^^^^^^^^^^^^^^^^^^^^^
Processing CSV DATA

Writing Results into Tables
vvvvvvvvvvvvvvvvvvvvvvvvv
"""
# %% Write to Excel https://stackoverflow.com/questions/32957441/putting-many-python-pandas-dataframes-to-one-excel-worksheet
int_name = list(set(int_name))
int_time = list(set(int_time))
int_scenario = list(set(int_scenario))
int_key = list(set(int_key))
int_node = list(set(int_node))
int_nodekey = list(set(int_nodekey))
dict_acrynm = {'Highway': 'Hwy',
               'Road': 'Rd',
               'Boulevard':'Blvd',
               'Street': 'St',
               'Lane': 'Ln',
               'Avenue': 'Ave',
               'Circle': 'Circ',
               'Court': 'Ct',
               'Crescent': 'Cres',
               'Drive': 'Dr',
               'Expressway': 'Exp',
               'Parkway': 'Pkwy',
               'Freeway': 'Fwy',
               'Gardens': 'Gdns',
               'Heights': 'Ht',
               'Hollow': 'Hllw',
               'Wonder': 'Wndr',
               'Land': 'Lnd',
               'West': 'W',
               'North': 'N',
               'South': 'S',
               'East': 'E',
               'County': 'Cnty',
               '/': '',
               'Dominion': 'Dmn',
               'Highschool': 'HS',
               'High School': 'HS',
               'Orangeville': 'OM',
               'Hansen': 'Hnsn',
               'Credit Creek Village': 'CC'}

# =============================================================================
# # shorten names - by node # 
# =============================================================================
""" October 28, 2020 
https://stackoverflow.com/questions/32957441/putting-many-python-pandas-dataframes-to-one-excel-worksheet
"""
int_name_xcel = []
int_name2 = []
for i in range(len(int_name)):
    int_name2.append(int_nodekey[i][1])
for i in range(len(int_name)):
    name = int_nodekey[i][0]
    int_name_xcel.append(name)
# write to excel - unfiltered
writer = pd.ExcelWriter('Results.xlsx',engine='xlsxwriter')
excel_row = [1]*len(int_name2)
for key, df in dict_table.items():
    index = int_name.index(key[0])
    v = key[4]
    temp = int_name_xcel.index(v)
    df.to_excel(writer,sheet_name = int_name_xcel[temp],startrow=excel_row[index], startcol=1)
    worksheet = writer.sheets[v]
    worksheet.write(excel_row[index]+1,0, key[1] + ' // ' + key[2])             # https://stackoverflow.com/questions/43537598/write-strings-text-and-pandas-dataframe-to-excel
    worksheet.write(0,0,key[0])                                                # added October 28, 2020
    excel_row[index] += len(df.index) +1                                        # add to its corresponding 1 in excel_row
writer.save()
#run again for filtered results
if result_type == 'critical':
    writer = pd.ExcelWriter('Results - Filtered.xlsx', engine = 'xlsxwriter')
    excel_row = [1]*len(int_name2)
    for key, df in dict_table_filtered.items():
        if df.empty:                                                            # only writes to excel if there are critical movements
            pass
        else:
            index = int_name.index(key[0])
            v = key[4]
            temp = int_name_xcel.index(v)
            df.to_excel(writer,sheet_name = int_name_xcel[temp],startrow=excel_row[index], startcol=1)
            worksheet = writer.sheets[v]
            worksheet.write(excel_row[index]+1,0, key[1] + ' // ' + key[2])             
            worksheet.write(0,0,key[0])                                                
            excel_row[index] += len(df.index) +1                                       
    writer.save()
del temp
print('\nYour results have been printed to an Excel File')
input("Press 'Enter' to continue")

#%% Results to Table
from docx import Document
# getting bottom-right header from Synchro Prep.py . HEADERS HAS TO MATCH HEADERS IN CSV RESULTS FILE
scen_list = []
with open ('Report Headers.txt') as csvfile:
    csv_reader = csv.reader(csvfile, delimiter = '\t')
    for row in csv_reader:
        scen_list.append(row[2])
        print(row[1] + " - " + row[2] + " HAS BEEN APPENDED TO SCENARIO LIST")
# make unique list while preserving order : https://stackoverflow.com/questions/480214/how-do-you-remove-duplicates-from-a-list-whilst-preserving-order
def UniqueList(seq):
    seen = set()
    seen_add = seen.add
    return [x for x in seq if not (x in seen or seen_add(x))]
scen_list = UniqueList(scen_list)
#templates
doc_sig = Document(path + '\Signalized.docx')
# takes special character from movement column << allows 'textwrapping'. uniqie to microsft?? each cell in the table shows a 'space' between each movement result, but this 'space' is not an ordinary space
special_char = doc_sig.tables[0].cell(2,1).text[5]
del doc_sig
# =============================================================================
# The document.save("file.docx") function will not replace the existing file until it is opened as document object (document = Document("file.docx"))
# https://stackoverflow.com/questions/56345124/python-docx-module-not-overwriting-same-file-name
# =============================================================================
dict_docx = {}
for key in dict_table:
    int_id = key[4]
    if key[3] == 'Lanes, Volumes, Timings':
        #this works... copy.deepcopy does not work....
        dict_docx[int_id] = Document(path + '\Signalized.docx')
    elif key[3] == 'HCM Unsignalized Intersection Capacity Analysis':
        # check if AWSC or TWSC
        if 'Overall' in list(dict_table[key].index): 
            dict_docx[int_id] = Document(path + '\AWSC.docx')
        else:
            dict_docx[int_id] = Document(path + '\TWSC.docx')
for key in dict_table:
    # strings of results
    a = ''
    b = ''
    c = ''
    d = ''
    e = ''
    """"important naming convention"""    
    #provided scenarios are named (20xx [Existing.../Future Background.../Total Future...])
    scenario = key[2]
    int_name = key[0]
    int_id = key[4]  
    if key[3] == 'Lanes, Volumes, Timings':
        # movement
        movement = list(dict_table[key].index)[:-1]
        # volume to capacity
        vc = list(dict_table[key]['v/c Ratio'])[:-1]
        # level of service
        los = list(dict_table[key]['LOS'])[:-1]
        # delay
        delay = list(dict_table[key]['Total Delay'])[:-1]
        # queue
        queue = list(dict_table[key]['Queue Length 95th (m)'])[:-1]
        # overall LOS
        overall_LOS = dict_table[key].loc['Overall','LOS']
        # overall delay
        overall_delay = dict_table[key].loc['Overall', 'Total Delay']       
        #counting rows
        row_count = 0
        for table in dict_docx[int_id].tables:
            for row in table.rows:
                row_count += 1        
        # replace scenario column with scenario names from 'Report Headers.txt'    
        for i, scen in zip(range(2, row_count+2, 2),scen_list):
            dict_docx[int_id].tables[0].cell(i,0).text = scen
        # for signalized and AWSC : 2 rows for headers. divide by 2 beause in signalized and AWSC each secnairo has two rows (one for results, one for overall)
        num_scenario = int((row_count -2 )/ 2)                                          
        result_row = []
        overall_row = []
        for v in range(num_scenario):
            result_row.append(2*(v+1))
            overall_row.append(2*(v+1)+1)
        # finds first instance of match
        for i in range(2,row_count):
            if scenario in dict_docx[int_id].tables[0].cell(i,0).text:
                # in python range
                desired_row = i
                break
        for j in range(len(vc)):
            if j != len(vc) - 1:
                a += movement[j] + special_char
                b += str(vc[j]) + special_char
                c += los[j] + special_char
                d += str(delay[j]) + special_char
                e += str(queue[j]) + special_char    
            elif j == len(vc) -1:
                a += movement[j]
                b += str(vc[j])
                c += los[j]
                d += str(delay[j])
                e += str(queue[j])             
# =============================================================================
#         Identify peak hour names
# =============================================================================
        # default - used in typical AM/PM/Midday projects
        if 'AM' in key[1]:
            y = 0
        elif 'PM' in key[1]:
            y = 4
        elif 'Saturday' or 'Midday' in key[1]:
            y = 8
        # special cases - for example, used in proj# 20-2907 Brethren Church
        # if 'Weekday' in key[1]:
        #     y = 0
        # elif 'Saturday' in key[1]:
        #     y = 4
        # elif 'Sunday' in key[1]:
        #     y = 8
        
        # # used for wasaga beach wendy
        # if 'Midday' in key[1]:
        #     y = 0
        # elif 'PM' in key[1]:
        #     y = 4
        
        dict_docx[int_id].tables[0].cell(desired_row,1).text = a   # movement
        dict_docx[int_id].tables[0].cell(desired_row,2+y).text = b   #v/c
        dict_docx[int_id].tables[0].cell(desired_row,3+y).text = c  # LOS
        dict_docx[int_id].tables[0].cell(desired_row,4+y).text = d   # delay
        dict_docx[int_id].tables[0].cell(desired_row,5+y).text = e   #queue
        dict_docx[int_id].tables[0].cell(desired_row+1,3+y).text = overall_LOS
        dict_docx[int_id].tables[0].cell(desired_row+1,4+y).text = overall_delay    
        dict_docx[int_id].add_paragraph(int_name)
        # dict_docx[int_id].save(path + '\Signalized - ' + int_id + '.docx')
        dict_docx[int_id].save(path + '\Signalized - ' + int_id + ' - ' + Name_Shorten(int_name,dict_acrynm) + '.docx')
    elif key[3] == 'HCM Unsignalized Intersection Capacity Analysis':
        # check if AWSC or TWSC
        if 'Overall' in list(dict_table[key].index):
            # movement
            movement = list(dict_table[key].index)[:-1]
            # volume to capacity
            vc = list(dict_table[key]['Degree Utilization, x'])[:-1]
            # level of service
            los = list(dict_table[key]['Approach LOS'])[:-1]
            # delay
            delay = list(dict_table[key]['Control Delay (s)'])[:-1]
            # overall LOS
            overall_LOS = dict_table[key].loc['Overall','Approach LOS']
            # overall delay
            overall_delay = dict_table[key].loc['Overall', 'Control Delay (s)']
            #counting rows
            row_count = 0
            for table in dict_docx[int_id].tables:
                for row in table.rows:
                    row_count += 1      
            # replace scenario column with scenario names from 'Report Headers.txt'    
            for i, scen in zip(range(2, row_count+2, 2),scen_list):
                dict_docx[int_id].tables[0].cell(i,0).text = scen
            # finds first instance of match
            for i in range(2,row_count):                
                if scenario in dict_docx[int_id].tables[0].cell(i,0).text:
                    # in python range
                    desired_row = i
                    # print(desired_row)
                    break                   
            # for signalized and AWSC
            num_scenario = int((row_count -2 )/ 2)                                          # 2 rows for headers. divide by 2 beause in signalized and AWSC each secnairo has two rows (one for results, one for overall)
            result_row = []
            overall_row = []
            for v in range(num_scenario):
                result_row.append(2*(v+1))
                overall_row.append(2*(v+1)+1)
            for j in range(len(vc)):
                if j != len(vc) - 1:
                    a += movement[j] + special_char
                    b += str(vc[j]) + special_char
                    c += los[j] + special_char
                    d += str(delay[j]) + special_char  
                elif j == len(vc) -1:
                    a += movement[j]
                    b += str(vc[j])
                    c += los[j]
                    d += str(delay[j])
# =============================================================================
#             Identify peak hour names
# =============================================================================
            #  default used in typical AM/PM/Midday projects
            if 'AM' in key[1]:
                y = 0
            elif 'PM' in key[1]:
                y = 3
            elif 'Saturday' or 'Midday' in key[1]:
                y = 6
            # Special cases - for example, used in proj# 20-2907 Brethren Church
            # if 'Weekday' in key[1]:
            #     y = 0
            # elif 'Saturday' in key[1]:
            #     y = 3
            # elif 'Sunday' in key[1]:
            #     y = 6            
            
            # # used in wassaga beach wendy
            # if 'Midday' in key[1]:
            #     y = 0
            # elif 'PM' in key [1]:
            #     y = 3
                
            dict_docx[int_id].tables[0].cell(desired_row,1).text = a   # movement
            dict_docx[int_id].tables[0].cell(desired_row,2+y).text = b   #v/c
            dict_docx[int_id].tables[0].cell(desired_row,3+y).text = c  #LOS
            dict_docx[int_id].tables[0].cell(desired_row,4+y).text = d   # delay
            dict_docx[int_id].tables[0].cell(desired_row+1,3+y).text = overall_LOS
            dict_docx[int_id].tables[0].cell(desired_row+1,4+y).text = overall_delay          
            dict_docx[int_id].add_paragraph(int_name)
            # dict_docx[int_id].save(path + '\AWSC - ' + int_id + '.docx')
            dict_docx[int_id].save(path + '\AWSC - ' + int_id + ' - ' + Name_Shorten(int_name,dict_acrynm) + '.docx')
        else: 
            # movement
            movement = list(dict_table[key].index)
            # volume to capacity
            vc = list(dict_table[key]['Volume to Capacity'])
            # level of service
            los = list(dict_table[key]['Lane LOS'])
            # delay
            delay = list(dict_table[key]['Control Delay (s)'])
            # queue
            queue = list(dict_table[key]['Queue Length 95th (m)'])
            #counting rows
            row_count = 0
            for table in dict_docx[int_id].tables:
                for row in table.rows:
                    row_count += 1                        
            # replace scenario column with scenario names from 'Report Headers.txt'    
            for i, scen in zip(range(2, row_count),scen_list):                
                dict_docx[int_id].tables[0].cell(i,0).text = scen        
            # finds first instance of match
            for i in range(2,row_count):
                if scenario in dict_docx[int_id].tables[0].cell(i,0).text:
                    # in python range
                    desired_row = i
                    break
            #making string of results
            for j in range(len(vc)):
                if j != len(vc) - 1:
                    a += movement[j] + special_char
                    b += str(vc[j]) + special_char
                    c += los[j] + special_char
                    d += str(delay[j]) + special_char
                    e += str(queue[j]) + special_char    
                elif j == len(vc) -1:
                    a += movement[j]
                    b += str(vc[j])
                    c += los[j]
                    d += str(delay[j])
                    e += str(queue[j])
# =============================================================================
#               Identify Peak hour names
# =============================================================================
            # Typical AM/PM/Midday projects
            if 'AM' in key[1]:
                y = 0
            elif 'PM' in key[1]:
                y = 4
            elif 'Saturday' or 'Midday' in key[1]:
                y = 8
            # special cases - for example, used in proj# 20-2907 Brethren Church
            # if 'Weekday' in key[1]:
            #     y = 0
            # elif 'Saturday' in key[1]:
            #     y = 4
            # elif 'Sunday' in key[1]:
            #     y = 8
            
            # # used for Wasaga Beach Wendy
            # if 'Midday' in key[1]:
            #     y = 0
            # elif 'PM' in key[1]:
            #     y = 4
            dict_docx[int_id].tables[0].cell(desired_row,1).text = a   # movement
            dict_docx[int_id].tables[0].cell(desired_row,2+y).text = b   #v/c
            dict_docx[int_id].tables[0].cell(desired_row,3+y).text = c  #LOS
            dict_docx[int_id].tables[0].cell(desired_row,4+y).text = d   # delay
            dict_docx[int_id].tables[0].cell(desired_row,5+y).text = e   #queue
            dict_docx[int_id].add_paragraph(int_name)
            # dict_docx[int_id].save(path + '\TWSC - ' + int_id + '.docx')
            dict_docx[int_id].save(path + '\TWSC - ' + int_id + ' - ' + Name_Shorten(int_name,dict_acrynm) + '.docx')

#%% replace scenario names from Synchro Header to scenario name for report
# =============================================================================
# for default AM/PM/Mid
# # =============================================================================
def RenameScenario(string):
    global year, condition
    if '20' in scen_name.split(' ')[0]:
        year = scen_name.split(' ')[0]
    else:
        year = ''
    if 'Existing' in scen_name:
        condition = 'Existing'
    elif ('Future Background' in scen_name) or ('FB' in scen_name):
        condition = 'Future Background'
    elif ('Total Future' in scen_name) or ('TF' in scen_name):
        condition = 'Total Future'
    return year, condition
for docxfile in sorted(glob.glob("*.docx")):
    if docxfile == 'Signalized.docx':
        next
    elif docxfile == 'AWSC.docx':
        next
    elif docxfile == 'TWSC.docx':
        next
    else:
        if 'AWSC' in docxfile:
            doc_AWSC_final = Document(path + '\\' + docxfile)
            row_count = 0
            for table in doc_AWSC_final.tables:
                for row in table.rows:
                    row_count += 1      
            for i in range(2,row_count,2):
                scen_name = doc_AWSC_final.tables[0].cell(i,0).text
                # print(scen_name)
                year, condition = RenameScenario(scen_name)
                doc_AWSC_final.tables[0].cell(i,0).text = year + ' ' + condition
            doc_AWSC_final.save(path + '\\' + docxfile)
        elif "TWSC" in docxfile:
            doc_TWSC_final = Document(path + '\\' + docxfile)
            row_count = 0
            for table in doc_TWSC_final.tables:
                for row in table.rows:
                    row_count += 1      
            for i in range(2,row_count):
                scen_name = doc_TWSC_final.tables[0].cell(i,0).text
                year, condition = RenameScenario(scen_name)
                doc_TWSC_final.tables[0].cell(i,0).text = year + ' ' + condition
            doc_TWSC_final.save(path + '\\' + docxfile) 
        elif 'Signalized' in docxfile:
            doc_sig_final = Document(path + '\\' + docxfile)
            row_count = 0
            for table in doc_sig_final.tables:
                for row in table.rows:
                    row_count += 1      
            for i in range(2,row_count,2):
                scen_name = doc_sig_final.tables[0].cell(i,0).text
                year, condition = RenameScenario(scen_name)
                doc_sig_final.tables[0].cell(i,0).text = year + ' ' + condition
            doc_sig_final.save(path + '\\' + docxfile)
        else:
            # print('file unknown')
            next
# https://stackoverflow.com/questions/45853595/spyder-clear-variable-explorer-along-with-variables-from-memory#:~:text=Go%20to%20the%20IPython%20console,That's%20it.
def clear_all():
    """Clears all the variables from the workspace of the spyder application."""
    gl = globals().copy()
    for var in gl:
        if var[0] == '_': continue
        if 'func' in str(globals()[var]): continue
        if 'module' in str(globals()[var]): continue

        del globals()[var]
if __name__ == "__main__":
    clear_all()